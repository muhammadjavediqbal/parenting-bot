from operator import itemgetter
from langchain_core.prompts import ChatPromptTemplate
from langchain.prompts.prompt import PromptTemplate
import os
from langchain_pinecone import PineconeVectorStore
from langchain_openai import ChatOpenAI
from langchain_openai import OpenAIEmbeddings
from dotenv import load_dotenv
from langchain.retrievers.multi_query import MultiQueryRetriever
from langchain.retrievers import ContextualCompressionRetriever
from langchain.retrievers.document_compressors import  LLMChainFilter
from langchain.output_parsers.boolean import BooleanOutputParser
from langchain_core.runnables import RunnableParallel, RunnablePassthrough
from docx import Document as DocxDocument
from io import BytesIO
from langchain_community.docstore.document import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
import re

load_dotenv()
embeddings = OpenAIEmbeddings(model="text-embedding-3-small",openai_api_key=os.getenv("OPENAI_API_KEY"))

index="parenting-bot"


def format_response_bullet_points(content):
    # Split the content into parts: before bullets, bullets, and after bullets
    intro, *sections = re.split(r'(\d+\.\s)', content, maxsplit=1)
    
    # Handle the case where there are no numbered sections
    if not sections:
        return content.strip()

    # Process numbered sections
    numbered_sections = re.split(r'(\d+\.\s)', ''.join(sections))
    formatted_sections = []
    for i in range(1, len(numbered_sections), 2):
        number = numbered_sections[i].strip()
        text = numbered_sections[i + 1].strip()
        title, details = text.split(':', 1)
        formatted_sections.append(f"{number} {title.strip()}:\n   {details.strip()}")

    # Join the formatted parts
    formatted_bullets = "\n\n".join(formatted_sections)
    return f"{intro.strip()}\n\n{formatted_bullets}"


def remove_asterisks(response: str) -> str:
    return re.sub(r'\*', '', response)

def format_docs(docs):
    return "\n\n".join(doc.page_content for doc in docs)

def data_insert(documents,embeddings,index_name):
    try:
        doc_search=PineconeVectorStore.from_documents(
            documents,
            embedding=embeddings,
            index_name=index_name,
            # namespace="social-problems"
            )
        return doc_search
    except Exception as ex:
        return ex


def extract_docx_data(content):
    """Extract text from a DOCX file and return it in a structured format."""
    docs = []
    doc = DocxDocument(BytesIO(content))
    text_list=[]
    for paragraph in doc.paragraphs:
        text_list.append(paragraph.text)
        
    final_text=[" ".join(text_list)]
    chunks=doc_spliter(final_text,'docx')
    return chunks

def doc_spliter(pages,format):
    try:
        doc_list=[]
        textsplit = RecursiveCharacterTextSplitter(
                    separators=["\n\n", "\n",".",""],
                    chunk_size=1500, chunk_overlap=150,
                    length_function=len)
        for i,page in enumerate(pages):
            pg_splits = textsplit.split_text(page)

            for pg_sub_split in pg_splits:
                metadata = {'format':format,"page_no":i+1}
                doc_string = Document(page_content=pg_sub_split, metadata=metadata)
                doc_list.append(doc_string)
        return doc_list
    except:
        return None



def _combine_documents3(docs, document_separator="\n\n"):
    try:

        metadata_allowed_range = 3
        if docs:
            combined_list = [
            f"content:{doc.page_content} \n metadata:{doc.metadata if index < metadata_allowed_range else {}}"
            for index, doc in enumerate(docs)
            ]
            combined = document_separator.join(combined_list)
        else:
            combined = ""    
        return combined
    except Exception as ex:
        raise ex

def get_response(query1):
    try:
        prompt_str = """

You are Parenting Assistant, tasked with delivering responses in a formal, authoritative, and concise tone. Adhere to the following guidelines:

- If the query is unrelated to the context {context}, respond with: "This query is outside the scope of the provided context."
- Only process the query {question} if it directly aligns with the context {context}.
- Provide extremely concise, accurate answers for clarity and completeness. 2-3 sentences are sufficient.
- Structure responses with clear paragraph separation for improved readability.

query: {question}  
context: {context}
"""
        bot_prompt = ChatPromptTemplate.from_template(prompt_str)
        chat_llm_multiquery = ChatOpenAI(model='gpt-4o-mini',openai_api_key=os.getenv("OPENAI_API_KEY"),temperature=0)
        num_chunks=6
        
        try:
            vecdb = PineconeVectorStore.from_existing_index(index_name=index,
                                                        embedding=embeddings)
        except Exception as e:
            return False
        retriever = vecdb.as_retriever(search_type="similarity",
                                        search_kwargs={"k": num_chunks})
        
        retriever_mq = MultiQueryRetriever.from_llm(retriever=retriever, llm=chat_llm_multiquery, include_original=False)
        chat_llm_4o_mini = ChatOpenAI(model='gpt-4o-mini',openai_api_key=os.getenv("OPENAI_API_KEY"),temperature=0)
        reranker_str = """You are a grader assessing relevance of a retrieved document to a user question. Your goal is to filter out erroneous retrievals and test does not need to be stringent.\n 
            If the document contains keyword(s) or semantic meaning related to the user question, return 'YES', otherwise return 'NO'. \n 
            Give a binary score 'YES' or 'NO' score to indicate whether the document is relevant to the question.
            
            > Question: {question}
            > Context:
            >>>
            {context}
            >>>

            The output must be in the following format: 'YES'/'NO'
            """
        reranker_prompt = PromptTemplate(
            template=reranker_str,
            input_variables=["question", "context"],
            output_parser=BooleanOutputParser(),
        )
        compressor_llm = LLMChainFilter.from_llm(llm=chat_llm_4o_mini, prompt=reranker_prompt)
        rerank_retriever = ContextualCompressionRetriever(base_compressor=compressor_llm,base_retriever=retriever_mq)
        setup_and_retrieval = RunnableParallel(
            {"context": rerank_retriever, "question": RunnablePassthrough()}
        )
        relevant_chunks=setup_and_retrieval.invoke(query1)
        if not relevant_chunks['context']:
            message="Sorry, it appears your query doesn't relate to Parenting."
            return message
        context=_combine_documents3(relevant_chunks['context'])
        chat_llm = ChatOpenAI(model="gpt-4o-mini", api_key=os.getenv("OPENAI_API_KEY"))
        
        query_fetcher= itemgetter("question")
        context_fetcher=itemgetter("context")
        setup={"question":query_fetcher,"context":context_fetcher}
        bot_chain = (setup | bot_prompt | chat_llm)
        response=bot_chain.invoke({"question":query1,"context":context}).content
        final_response = remove_asterisks(response) 
        return final_response
    except Exception as ex:
        return str(ex)